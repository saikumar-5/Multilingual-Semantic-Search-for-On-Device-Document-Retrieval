"""
DOCX Parser using python-docx.
Extracts text from Word documents including paragraphs and tables.

CO1 Alignment: Document Representation - extracting text from DOCX documents.
"""

from pathlib import Path
from typing import Optional
import logging

logger = logging.getLogger(__name__)


class DocxParser:
    """Extract text from .docx files."""

    def extract(self, file_path: Path) -> dict:
        """
        Extract text from a DOCX file.

        Returns:
            dict with keys: 'text', 'metadata', 'file_path', 'file_type'
        """
        from docx import Document

        file_path = Path(file_path)
        result = {
            "file_path": str(file_path),
            "file_name": file_path.name,
            "file_type": "docx",
            "text": "",
            "metadata": {},
        }

        try:
            doc = Document(str(file_path))

            # Extract paragraphs
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

            # Extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        table_texts.append(row_text)

            all_text = paragraphs + table_texts
            result["text"] = "\n".join(all_text)

            # Metadata from core properties
            if doc.core_properties:
                result["metadata"] = {
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                }

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            result["text"] = ""

        return result
